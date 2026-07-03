"""DOCX document processor using python-docx."""

import asyncio
import logging
from datetime import datetime

logger = logging.getLogger(__name__)


def _load_docx():
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX processing. Install with: pip install python-docx"
        ) from None
    return Document


class DOCXProcessor:
    """Extract text and metadata from DOCX files."""

    async def extract_text(self, file_path: str) -> str:
        """Extract text from a DOCX file.

        Walks paragraphs, tables, and section headers/footers. Forms and
        reports often put content inside tables, so extracting body paragraphs
        alone misses it. Text boxes, inline shapes, and footnotes are not
        covered (they need XML traversal python-docx doesn't expose directly).
        """
        Document = _load_docx()

        def _extract():
            doc = Document(file_path)
            text_parts: list[str] = []

            # Body paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)

            # Body tables — walk rows x cells, joining cell text with tabs.
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        text_parts.append("\t".join(cells))

            # Headers / footers (one per section) — often hold title/metadata.
            for section in doc.sections:
                for hf in (section.header, section.footer):
                    if hf is None:
                        continue
                    for paragraph in hf.paragraphs:
                        if paragraph.text:
                            text_parts.append(paragraph.text)

            return "\n".join(text_parts)

        return await asyncio.to_thread(_extract)

    async def extract_metadata(self, file_path: str) -> dict:
        """Best-effort metadata extraction — returns {} on any failure."""
        Document = _load_docx()

        def _extract():
            try:
                doc = Document(file_path)
                props = getattr(doc, "core_properties", None)
                if props is None:
                    return {}
                author = (getattr(props, "author", None) or "").strip() or None
                title = (getattr(props, "title", None) or "").strip() or None
                created = getattr(props, "created", None)
                authored_at = created if isinstance(created, datetime) else None
                return {"author": author, "doc_title": title, "authored_at": authored_at}
            except Exception as e:
                logger.debug("DOCX metadata extraction failed for %s: %s", file_path, e)
                return {}

        return await asyncio.to_thread(_extract)
